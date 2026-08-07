from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

def rtl(p):
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement('w:bidi')
    pPr.insert(0, b)

def para(doc, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT, color=None, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name = 'B Nazanin'
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    rtl(p)
    return p

def heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = h.add_run(text)
    r.font.name = 'B Nazanin'
    r.font.size = Pt(16 if level==1 else 13)
    r.font.bold = True
    if level==1: r.font.color.rgb = RGBColor(68,114,196)
    rtl(h)

def hline(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'12'); bot.set(qn('w:color'),'4472C4')
    pBdr.append(bot); pPr.append(pBdr)

def table(doc, rows_data, style='Light Grid Accent 1'):
    t = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    t.style = style
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = ''
            r = c.paragraphs[0].add_run(val)
            r.font.name = 'B Nazanin'
            r.font.bold = (i==0)
    return t

# ── جلد ──
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('گزارش پروژه PetroOpt AI'); r.font.name='B Nazanin'; r.font.size=Pt(24); r.font.bold=True; r.font.color.rgb=RGBColor(68,114,196)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('سیستم هوشمند بهینه‌سازی تولید پتروشیمی'); r2.font.name='B Nazanin'; r2.font.size=Pt(16); r2.font.color.rgb=RGBColor(68,114,196)
doc.add_paragraph()
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ['نویسنده: [نام دانشجو]','رشته: مهندسی صنایع','دانشگاه صنعتی شریف','تاریخ: ۱۴۰۵']:
    r3 = p3.add_run(line+'\n'); r3.font.name='B Nazanin'; r3.font.size=Pt(13)
doc.add_page_break()

# ── فهرست ──
heading(doc,'فهرست مطالب')
for item in ['۱. معرفی پروژه','۲. ساختار پروژه','۳. نحوه کارکرد سیستم',
             '۴. ماژول برنامه‌ریزی خطی (LP)','۵. ماژول الگوریتم ژنتیک (GA)',
             '۶. ماژول دستیار گفت‌وگویی','۷. رابط کاربری',
             '۸. اعتبارسنجی ورودی‌ها','۹. خروجی‌های سیستم',
             '۱۰. نصب و اجرا','۱۱. جمع‌بندی']:
    para(doc, item)
doc.add_page_break()

# ── ۱. معرفی ──
heading(doc,'۱. معرفی پروژه'); hline(doc)
para(doc,'PetroOpt AI یک سامانه هوشمند پشتیبان تصمیم برای برنامه‌ریزی تولید پتروشیمی است که سه فناوری اصلی را ترکیب می‌کند:')
for t_ in ['• برنامه‌ریزی خطی (LP): بهینه‌سازی برنامه تولید و بیشینه‌سازی سود',
           '• الگوریتم ژنتیک (GA): زمان‌بندی بهینه تعمیرات واحدها',
           '• دستیار گفت‌وگویی: تحلیل نتایج و پاسخ به پرسش‌های تحلیلی']:
    para(doc, t_)

# ── ۲. ساختار ──
heading(doc,'۲. ساختار پروژه'); hline(doc)
for line in ['• app.py: نقطه ورود برنامه و رابط کاربری Streamlit',
             '• models/: مدل‌های داده (refinery, maintenance, result)',
             '• lp/: ماژول برنامه‌ریزی خطی (optimizer, model, constraints)',
             '• ga/: ماژول الگوریتم ژنتیک (genetic, chromosome, crossover, mutation, selection, population, fitness)',
             '• chatbot/: دستیار گفت‌وگویی (assistant, conversation, memory)',
             '• utils/: ابزارهای کمکی (charts, validator)']:
    para(doc, line)
doc.add_page_break()

# ── ۳. نحوه کارکرد ──
heading(doc,'۳. نحوه کارکرد سیستم'); hline(doc)
heading(doc,'۳.۱ جریان اجرایی کلی', 2)
for line in ['۱. ورود داده‌ها توسط کاربر در رابط Streamlit',
             '۲. اعتبارسنجی داده‌ها با InputValidator',
             '۳. بهینه‌سازی تولید با برنامه‌ریزی خطی (PuLP/CBC)',
             '۴. زمان‌بندی تعمیرات با الگوریتم ژنتیک',
             '۵. ذخیره نتایج در session_state',
             '۶. نمایش داشبورد (سود، تولید، منابع، نمودار)',
             '۷. پاسخ به پرسش‌های تحلیلی از طریق دستیار']:
    para(doc, line)
heading(doc,'۳.۲ ورودی‌های سیستم', 2)
for line in ['• اطلاعات عمومی پالایشگاه: کل خوراک (تن)، کل انرژی، تعداد واحدها',
             '• اطلاعات هر واحد: نام، ظرفیت، سود/تن، مصرف خوراک، مصرف انرژی',
             '• اطلاعات تعمیرات: نام واحد، مدت (روز)، اولویت (۱-۱۰)، بازه زمانی مجاز']:
    para(doc, line)
doc.add_page_break()

# ── ۴. LP ──
heading(doc,'۴. ماژول برنامه‌ریزی خطی (LP)'); hline(doc)
heading(doc,'۴.۱ مدل ریاضی', 2)
para(doc,'متغیرهای تصمیم: برای هر واحد i، مقدار تولید x_i تعریف می‌شود:')
para(doc,'0 ≤ x_i ≤ capacity_i', align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,'تابع هدف (بیشینه‌سازی سود):', bold=True)
para(doc,'max Z = Σ ( x_i × profit_i )', align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,'محدودیت خوراک:', bold=True)
para(doc,'Σ ( x_i × feed_i ) ≤ total_feed', align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,'محدودیت انرژی:', bold=True)
para(doc,'Σ ( x_i × energy_i ) ≤ total_energy', align=WD_ALIGN_PARAGRAPH.CENTER)
heading(doc,'۴.۲ خروجی بهینه‌ساز', 2)
for line in ['• وضعیت حل: Optimal / Infeasible',
             '• حداکثر سود قابل دستیابی',
             '• برنامه تولید هر واحد (تن)',
             '• منابع مصرف‌شده و باقی‌مانده']:
    para(doc, line)
doc.add_page_break()

# ── ۵. GA ──
heading(doc,'۵. ماژول الگوریتم ژنتیک (GA)'); hline(doc)
para(doc,'هدف: زمان‌بندی بهینه تعمیرات با کمینه‌سازی تداخل و رعایت محدودیت‌های زمانی.')
heading(doc,'۵.۱ پارامترهای الگوریتم', 2)
table(doc,[('پارامتر','مقدار'),('اندازه جمعیت','۲۰ کروموزوم'),('تعداد نسل‌ها','۵۰'),
           ('روش انتخاب','Tournament Selection'),('روش ترکیب','Order Crossover (OX)'),
           ('روش جهش','Swap Mutation')])
doc.add_paragraph()
heading(doc,'۵.۲ جریان الگوریتم', 2)
for line in ['۱. ایجاد جمعیت اولیه (تصادفی)',
             '۲. ارزیابی Fitness هر کروموزوم',
             '۳. حلقه ۵۰ نسل: انتخاب → ترکیب → جهش → ارزیابی → جایگزینی',
             '۴. نگاشت بهترین کروموزوم به برنامه تعمیرات',
             'نکته: با یک وظیفه تعمیراتی، GA اجرا نمی‌شود و earliest_start استفاده می‌شود.']:
    para(doc, line)
doc.add_page_break()

# ── ۶. دستیار ──
heading(doc,'۶. ماژول دستیار گفت‌وگویی'); hline(doc)
para(doc,'معماری: Rule-based Intent Detection با تطبیق کلیدواژه و حافظه مکالمه (last_intent).')
heading(doc,'۶.۱ نیت‌های پشتیبانی‌شده', 2)
table(doc,[('نیت','توضیح'),('greeting','خوش‌آمدگویی'),('profit','تحلیل سود'),
           ('feed','تحلیل خوراک'),('energy','تحلیل انرژی'),('production','تحلیل تولید'),
           ('maintenance','تحلیل تعمیرات'),('recommendation','پیشنهادهای بهینه‌سازی'),
           ('summary','خلاصه نتایج'),('compare','مقایسه واحدها'),('bottleneck','شناسایی گلوگاه')])
doc.add_page_break()

# ── ۷. رابط کاربری ──
heading(doc,'۷. رابط کاربری (Streamlit)'); hline(doc)
for line in ['داشبورد نتایج شامل:',
             '• شاخص‌های اصلی: وضعیت حل، حداکثر سود، خوراک و انرژی مصرف‌شده',
             '• برنامه تولید: مقدار تولید هر واحد (تن)',
             '• جدول تولید و نمودار میله‌ای (ChartGenerator)',
             '• برنامه تعمیرات: روز تعمیر هر واحد',
             '• بخش دستیار: ورود پرسش و دریافت پاسخ تحلیلی']:
    para(doc, line)

# ── ۸. اعتبارسنجی ──
heading(doc,'۸. اعتبارسنجی ورودی‌ها'); hline(doc)
for line in ['validate_refinery(refinery):',
             '  • مقادیر خوراک و انرژی باید مثبت باشند',
             '  • حداقل یک واحد تولیدی باید وجود داشته باشد',
             '  • پارامترهای هر واحد باید معتبر باشند',
             'validate_maintenance_tasks(tasks):',
             '  • اولویت در بازه ۱ تا ۱۰',
             '  • بازه زمانی تعمیرات معتبر باشد',
             'در صورت خطا: پیام نمایش داده شده و اجرا با st.stop() متوقف می‌شود.']:
    para(doc, line)
doc.add_page_break()

# ── ۹. خروجی‌ها ──
heading(doc,'۹. خروجی‌های سیستم'); hline(doc)
table(doc,[('خروجی','نوع','محل نمایش'),
           ('وضعیت حل مدل LP','متن','داشبورد'),
           ('حداکثر سود','عدد','داشبورد'),
           ('برنامه تولید هر واحد','جدول + نمودار','داشبورد'),
           ('خوراک و انرژی مصرف‌شده','عدد','داشبورد'),
           ('برنامه زمان‌بندی تعمیرات','جدول','داشبورد'),
           ('پاسخ‌های تحلیلی','متن','بخش دستیار')])
doc.add_paragraph()

# ── ۱۰. نصب ──
heading(doc,'۱۰. نصب و اجرا'); hline(doc)
para(doc,'فناوری‌های اصلی: Python, Streamlit, PuLP, Pandas, Matplotlib')
para(doc,'نصب وابستگی‌ها:', bold=True)
p_c = doc.add_paragraph(); r_c = p_c.add_run('pip install -r requirements.txt'); r_c.font.name='Courier New'; r_c.font.size=Pt(10)
para(doc,'اجرای برنامه:', bold=True)
p_c2 = doc.add_paragraph(); r_c2 = p_c2.add_run('streamlit run app.py'); r_c2.font.name='Courier New'; r_c2.font.size=Pt(10)
doc.add_page_break()

# ── ۱۱. جمع‌بندی ──
heading(doc,'۱۱. جمع‌بندی'); hline(doc)
para(doc,'PetroOpt AI با ترکیب بهینه‌سازی ریاضی (LP)، محاسبات تکاملی (GA) و پردازش زبان طبیعی مبتنی بر قواعد، یک ابزار کامل پشتیبان تصمیم برای مدیران پتروشیمی فراهم می‌کند.')
for line in ['• معماری ماژولار: توسعه مستقل هر بخش',
             '• رابط کاربری ساده: ورود آسان داده و مشاهده نتایج',
             '• تحلیل هوشمند: پاسخ به پرسش‌های تحلیلی',
             '• کاربردها: برنامه‌ریزی تولید، بهینه‌سازی منابع، زمان‌بندی تعمیرات']:
    para(doc, line)
doc.add_paragraph()
p_f = doc.add_paragraph(); p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_f.add_run('این گزارش بر اساس بررسی مستقیم کدهای پروژه تهیه شده است.')
r_f.font.name='B Nazanin'; r_f.font.size=Pt(10); r_f.font.italic=True; r_f.font.color.rgb=RGBColor(128,128,128)

doc.save('PetroOpt_AI_Report.docx')
print('Done: PetroOpt_AI_Report.docx')
